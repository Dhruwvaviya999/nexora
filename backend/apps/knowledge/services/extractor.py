"""
Text extraction from uploaded documents.

Supports PDF, DOCX, TXT and Markdown. Heavy parsers (pypdf, python-docx) are
imported lazily so they're only required when a matching file is processed.
"""

from __future__ import annotations

import os

# Extensions we can extract text from. Anything else is skipped (no embeddings).
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class UnsupportedDocumentError(Exception):
    """Raised when a document's type cannot be turned into text."""


def is_supported(filename: str, content_type: str = "") -> bool:
    return _extension(filename) in SUPPORTED_EXTENSIONS


def _extension(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def extract_text(file, filename: str, content_type: str = "") -> str:
    """Return the plain text of ``file`` (an open, readable file-like object)."""
    ext = _extension(filename)
    if ext == ".pdf":
        return _extract_pdf(file)
    if ext == ".docx":
        return _extract_docx(file)
    if ext in {".txt", ".md", ".markdown"}:
        return _extract_text(file)
    raise UnsupportedDocumentError(
        f"Unsupported document type: {ext or content_type or 'unknown'}"
    )


def _extract_pdf(file) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(file) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Include table cell text — meeting notes often live in tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(file) -> str:
    raw = file.read()
    if isinstance(raw, bytes):
        # Be forgiving about encodings rather than failing the whole job.
        return raw.decode("utf-8", errors="replace")
    return raw
